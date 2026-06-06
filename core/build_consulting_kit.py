"""
Generates Noah_Hawkes_Consulting_Kit.docx
Ready-to-paste copy for LinkedIn, Upwork, Clarity.fm, and cold outreach.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("G:/My Drive/HawkesNest LLC/Noah.AI/Noah_Hawkes_Consulting_Kit.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)


# ── Helpers ────────────────────────────────────────────────────────────────────
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
ACCENT = RGBColor(0x00, 0x6E, 0xC0)   # LinkedIn blue
GRAY   = RGBColor(0x55, 0x55, 0x55)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '006EC0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def divider():
    doc.add_paragraph()

def label_value(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(10)
t = title.add_run("NOAH HAWKES — CONSULTING LAUNCH KIT")
t.bold = True
t.font.size = Pt(16)
t.font.color.rgb = DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
s = sub.add_run("Ready-to-paste copy for LinkedIn · Upwork · Clarity.fm · Cold Outreach")
s.font.size = Pt(10)
s.font.color.rgb = GRAY
s.italic = True

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(20)
c = contact.add_run("noahhawkes@gmail.com  ·  801-900-9049  ·  Georgetown, TX  ·  linkedin.com/in/noahhawkes")
c.font.size = Pt(9.5)
c.font.color.rgb = ACCENT


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — LINKEDIN
# ══════════════════════════════════════════════════════════════════════════════
h1("SECTION 1 — LINKEDIN")

h2("Headline (copy into LinkedIn > Edit Profile > Headline)")
body('Operations Fixer & SOP Architect | I walk into broken companies and fix them | MBA | Trucking · Industrial · Small Business')

divider()

h2("About Section (copy into LinkedIn > Edit Profile > About)")
body(
    "I fix what's broken.\n\n"
    "Companies hire me when revenue is bleeding, teams are fractured, safety ratings are in the red, "
    "or operations have quietly become a liability. I walk in, diagnose fast, and get to work.\n\n"
    "My background spans industrial sales (Samco Technologies), transportation operations and dispatch "
    "(Hawkes Trucking), and cross-functional leadership across operations, compliance, and culture.\n\n"
    "What I actually do:\n"
)
bullet("Turnaround ops — rebuilding trust, cutting waste, restoring performance in 60–90 days")
bullet("SOP development — documenting processes so businesses stop depending on one person knowing everything")
bullet("Safety & compliance — OSHA, FMCSA, EPA risk reduction for trucking and industrial operations")
bullet("Team culture — transforming toxic environments into high-performing teams")
bullet("Revenue recovery — identifying and closing the gaps where money is walking out the door")
body(
    "\nI've led United Way projects that delivered 125% sales increases, chaired committees at the "
    "Minnesota State Chamber of Commerce, earned an MBA, and carry the rank of Eagle Scout and Master Freemason. "
    "I don't quit, and I don't leave a mess behind.\n\n"
    "If your business has a problem that needs someone who's seen it before and knows how to solve it — "
    "let's talk.\n\n"
    "📧 noahhawkes@gmail.com  |  📞 801-900-9049"
)

divider()

h2("Featured Section Post (pin this as your first post)")
body(
    '"I specialize in one thing: walking into a broken operation and fixing it.\n\n'
    "Toxic culture. Lost revenue. Safety violations. OSHA flags. Teams that stopped trusting each other.\n\n"
    "I've done it in industrial sales, trucking operations, and small business environments. "
    "I diagnose the real problem — not the symptom — and I build the systems and culture to make sure it doesn't come back.\n\n"
    "MBA. Eagle Scout. Master Freemason. Former dispatch and ops.\n\n"
    "If you're dealing with a problem in your company that feels unsolvable — it probably isn't. "
    'Send me a message."'
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — UPWORK
# ══════════════════════════════════════════════════════════════════════════════
h1("SECTION 2 — UPWORK")

body("Go to upwork.com → Create Account → Professional profile. Use the copy below.", italic=True, color=GRAY)

h2("Profile Title")
body("Operations & Process Consultant | Business Turnaround | SOP Development | FMCSA & OSHA Compliance")

divider()

h2("Professional Overview (the main bio field)")
body(
    "I fix operational problems that are costing your business money, people, or reputation.\n\n"
    "With an MBA and a career spanning industrial sales management, trucking and dispatch operations, "
    "and cross-functional leadership, I bring a rare combination: I can read a P&L, rewrite an SOP, "
    "rebuild a team, and navigate FMCSA or OSHA compliance — often in the same engagement.\n\n"
    "Specialties:\n"
)
bullet("Standard Operating Procedure (SOP) development — from scratch or audit-and-rebuild")
bullet("Operations audits — find what's broken, document it, fix it")
bullet("Transportation/trucking compliance — FMCSA, safety ratings, dispatch process")
bullet("Culture and team rebuilds — toxic team environments, high turnover, trust deficits")
bullet("Revenue recovery consulting — P&L analysis, process gaps, accountability systems")
body(
    "\nI've delivered a 125% sales increase through a community engagement project, "
    "led chamber of commerce policy committees, and built ops systems from the ground up. "
    "I bring real-world experience, not just frameworks.\n\n"
    "I work best on fixed-scope projects with clear deliverables. "
    "Initial audit packages start at $500. Full SOP builds and turnaround engagements from $1,500."
)

divider()

h2("Proposed Service Packages (set these as fixed-price contracts)")

h2("Package 1 — Operations Audit  |  $500  |  5–7 days")
bullet("Intake call (90 min) — you walk me through the operation")
bullet("Written audit report — what's broken, why, and ranked by impact")
bullet("Priority fix list with recommended next steps")
bullet("30-min debrief call")
body("Deliverable: PDF audit report + priority matrix", italic=True)

h2("Package 2 — SOP Build (single department)  |  $1,500  |  2 weeks")
bullet("Discovery interview with department lead")
bullet("Document existing process")
bullet("Write clean, executable SOPs in your format or a standard template")
bullet("Training outline for new hires")
body("Deliverable: 5–15 SOPs depending on complexity + onboarding guide", italic=True)

h2("Package 3 — Full Turnaround Engagement  |  $3,500–$7,500  |  60–90 days")
bullet("Full ops audit")
bullet("Culture and team assessment")
bullet("SOP development across all key departments")
bullet("Weekly check-ins, 90-day roadmap, implementation support")
body("Deliverable: Complete operations playbook + 90-day action plan", italic=True)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — CLARITY.FM
# ══════════════════════════════════════════════════════════════════════════════
h1("SECTION 3 — CLARITY.FM")

body("Go to clarity.fm → Sign up → Set your rate → Paste bio below.", italic=True, color=GRAY)
body("Suggested rate: $3.33/min ($200/hr). Start there, raise after 5 calls.")

divider()

h2("Clarity.fm Bio")
body(
    "Operations consultant and business fixer with an MBA and 15+ years across industrial sales, "
    "trucking/dispatch operations, and cross-functional leadership.\n\n"
    "I help business owners and operators with:\n"
)
bullet("Diagnosing why their operations are bleeding money or people")
bullet("Building SOPs and process documentation from scratch")
bullet("FMCSA and trucking compliance issues")
bullet("Turning around toxic team cultures")
bullet("Revenue recovery and accountability systems")
body(
    "\nFormer Industrial Sales Manager (Samco Technologies). "
    "United Way Project Leader. Minnesota Chamber of Commerce Committee Chair. "
    "Eagle Scout. Master Freemason.\n\n"
    "If you have 30 minutes and a real problem — I'll give you a real answer."
)

divider()

h2("Topics to select on Clarity.fm profile")
bullet("Operations Management")
bullet("Standard Operating Procedures")
bullet("Small Business Consulting")
bullet("Turnaround Management")
bullet("Sales Management")
bullet("Trucking / Transportation")
bullet("Team Management")
bullet("Business Process Improvement")


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — COLD OUTREACH EMAIL
# ══════════════════════════════════════════════════════════════════════════════
h1("SECTION 4 — COLD OUTREACH EMAIL")

body(
    "Target: small and mid-size business owners dealing with operational pain. "
    "Find them on LinkedIn (search: 'Operations Manager' or 'Owner' + your target industry). "
    "Send via LinkedIn InMail or direct email. Personalize the [brackets].",
    italic=True, color=GRAY
)

divider()

h2("Email Template A — The Fixer (general ops)")
body(
    "Subject: Quick question about [Company Name]\n\n"
    "Hi [First Name],\n\n"
    "I came across [Company Name] and noticed [something specific — growth, a job posting for ops roles, a review mentioning culture issues, etc.].\n\n"
    "I'm an operations consultant who specializes in fixing the things that quietly drain businesses — "
    "broken processes, team trust issues, compliance gaps, revenue leaks. "
    "I've done this in industrial, trucking, and small business environments.\n\n"
    "I'm not pitching a retainer. I offer a flat-fee operations audit ($500) that gives you a clear picture "
    "of what's costing you money and what to fix first.\n\n"
    "Would a 20-minute call make sense this week?\n\n"
    "Noah Hawkes\n"
    "noahhawkes@gmail.com | 801-900-9049\n"
    "linkedin.com/in/noahhawkes"
)

divider()

h2("Email Template B — SOP King (process/documentation)")
body(
    "Subject: Your team might be one resignation away from chaos\n\n"
    "Hi [First Name],\n\n"
    "Most small businesses have their processes living inside one or two people's heads. "
    "When those people leave — or just have a bad week — everything slows down or breaks.\n\n"
    "I build SOPs. Clean, executable, actually-used documentation that lets your team run without "
    "you holding everything together.\n\n"
    "Single department SOP build starts at $1,500 and takes two weeks. "
    "You end up with documented processes, a training guide, and a system that scales.\n\n"
    "Is this something [Company Name] is dealing with right now?\n\n"
    "Noah Hawkes\n"
    "noahhawkes@gmail.com | 801-900-9049\n"
    "linkedin.com/in/noahhawkes"
)

divider()

h2("Email Template C — Trucking / FMCSA (niche, highest rate)")
body(
    "Subject: FMCSA compliance question for [Company Name]\n\n"
    "Hi [First Name],\n\n"
    "I work with trucking and transportation operators on compliance and operations — "
    "specifically FMCSA safety ratings, dispatch process, and the operational gaps that get carriers "
    "into trouble before they see it coming.\n\n"
    "I noticed [Company Name] operates in [region/sector] and wanted to reach out. "
    "A lot of carriers I talk to are dealing with similar issues: driver turnover, "
    "safety scores trending the wrong direction, or dispatch processes that work until they don't.\n\n"
    "I do a flat-fee operations audit ($500) that gives you a clear picture of your exposure and "
    "what to fix first. No retainer, no long-term commitment.\n\n"
    "Worth a 20-minute call?\n\n"
    "Noah Hawkes\n"
    "noahhawkes@gmail.com | 801-900-9049\n"
    "linkedin.com/in/noahhawkes"
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — THIS WEEK'S ACTION PLAN
# ══════════════════════════════════════════════════════════════════════════════
h1("SECTION 5 — THIS WEEK (DO THESE IN ORDER)")

body("Each step takes under 30 minutes. Do all five by Friday.", italic=True, color=GRAY)

h2("Day 1 — LinkedIn (30 min)")
bullet("Update headline with the copy from Section 1")
bullet("Paste the About section")
bullet("Write and post the Featured Section post")
bullet("Connect with 10 people in operations, trucking, or small business ownership")

h2("Day 2 — Upwork (30 min)")
bullet("Create account at upwork.com (use noahhawkes@gmail.com)")
bullet("Paste title and bio from Section 2")
bullet("Set up the three service packages at the listed prices")
bullet("Upload your Fixer resume as a portfolio item")

h2("Day 3 — Clarity.fm (15 min)")
bullet("Create account at clarity.fm")
bullet("Paste bio, select topics from Section 3")
bullet("Set rate at $3.33/min")
bullet("Share your Clarity link in your LinkedIn bio")

h2("Day 4 — First outreach (45 min)")
bullet("Search LinkedIn for 'Operations Manager' OR 'Owner' in trucking, industrial, or manufacturing")
bullet("Find 10 people whose company looks like it has operational pain (job postings, reviews, growth signals)")
bullet("Send Template A or C from Section 4 — personalize the [brackets]")
bullet("Goal: 10 messages sent")

h2("Day 5 — Follow up + repeat")
bullet("Check replies, respond same day")
bullet("Send 10 more outreach messages")
bullet("If no Upwork clients in week 1: lower Package 1 to $250 to get first review, raise after")

divider()

body(
    "One audit client per week = $500–$2,000/week depending on package. "
    "Three clients converting to SOP builds = $4,500+ in 30 days. "
    "This is the path.",
    color=DARK
)


# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"Saved: {OUT}")
